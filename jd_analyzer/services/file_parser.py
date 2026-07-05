"""
File Parser Service
Handles extraction of raw text from PDF, DOCX, and plain text files.
"""

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class FileParserService:
    """
    Parses uploaded files and extracts raw text content.
    Supports: PDF (.pdf), Word documents (.docx), Plain text (.txt)
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

    def parse(self, file_obj, filename: str) -> str:
        """
        Entry point: detect file type and delegate to the appropriate parser.

        Args:
            file_obj: File-like object (werkzeug FileStorage or BytesIO).
            filename: Original filename (used to determine extension).

        Returns:
            Extracted raw text string.

        Raises:
            ValueError: If the file type is unsupported.
            RuntimeError: If extraction fails.
        """
        ext = self._get_extension(filename)
        file_bytes = self._read_bytes(file_obj)

        if ext == ".pdf":
            return self._parse_pdf(file_bytes, filename)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(file_bytes, filename)
        elif ext == ".txt":
            return self._parse_text(file_bytes, filename)
        else:
            raise ValueError(
                f"Unsupported file type '{ext}'. "
                f"Allowed types: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _get_extension(filename: str) -> str:
        import os
        return os.path.splitext(filename.lower())[1]

    @staticmethod
    def _read_bytes(file_obj) -> bytes:
        """Read bytes, resetting pointer if possible."""
        if hasattr(file_obj, "read"):
            data = file_obj.read()
            if hasattr(file_obj, "seek"):
                file_obj.seek(0)
            return data
        return bytes(file_obj)

    def _parse_pdf(self, file_bytes: bytes, filename: str) -> str:
        """Extract text from a PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise RuntimeError(
                "PyMuPDF is not installed. Run: pip install pymupdf"
            )

        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages = []
            for page in doc:
                text = page.get_text("text")
                if text.strip():
                    pages.append(text)
            doc.close()
            result = "\n".join(pages)
            if not result.strip():
                raise RuntimeError(
                    f"No readable text found in '{filename}'. "
                    "The PDF may be scanned/image-based."
                )
            logger.info(f"[FileParser] PDF '{filename}' extracted {len(result)} chars.")
            return result
        except Exception as e:
            raise RuntimeError(f"PDF parsing failed for '{filename}': {e}") from e

    def _parse_docx(self, file_bytes: bytes, filename: str) -> str:
        """Extract text from a DOCX file using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError(
                "python-docx is not installed. Run: pip install python-docx"
            )

        try:
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            result = "\n".join(paragraphs)
            if not result.strip():
                raise RuntimeError(f"No readable text found in '{filename}'.")
            logger.info(f"[FileParser] DOCX '{filename}' extracted {len(result)} chars.")
            return result
        except Exception as e:
            raise RuntimeError(f"DOCX parsing failed for '{filename}': {e}") from e

    @staticmethod
    def _parse_text(file_bytes: bytes, filename: str) -> str:
        """Decode plain text, trying common encodings."""
        for encoding in ("utf-8", "utf-16", "latin-1", "cp1252"):
            try:
                result = file_bytes.decode(encoding)
                if result.strip():
                    logger.info(
                        f"[FileParser] TXT '{filename}' decoded with {encoding}, "
                        f"{len(result)} chars."
                    )
                    return result
            except (UnicodeDecodeError, LookupError):
                continue
        raise RuntimeError(
            f"Could not decode '{filename}' with common encodings."
        )
