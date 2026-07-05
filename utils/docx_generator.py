from io import BytesIO

def generate_ats_docx(user_data, profile_data, projects, certs, template_id='tpl_prof_03'):
    """
    Generates an ATS-compliant DOCX file from structured resume data.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise RuntimeError(f"python-docx dependency missing. Cannot generate DOCX. {e}")

    doc = Document()

    # Title / Name
    name = user_data.get('name', 'Name Not Provided')
    h1 = doc.add_heading(name, level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact Info
    contact_info = f"{user_data.get('email', '')} | {profile_data.get('linkedin', '')} | {profile_data.get('portfolio', '')}"
    p = doc.add_paragraph(contact_info)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Professional Summary
    summary = profile_data.get('professional_summary', '')
    if summary:
        doc.add_heading('Professional Summary', level=2)
        doc.add_paragraph(summary)

    # Skills
    skills = profile_data.get('skills', '')
    if skills:
        doc.add_heading('Skills', level=2)
        doc.add_paragraph(skills)

    # Experience / Achievements (Since experience isn't fully structured, we use achievements)
    achievements = profile_data.get('achievements', '')
    if achievements:
        doc.add_heading('Experience & Achievements', level=2)
        for ach in achievements.split('\n'):
            if ach.strip():
                doc.add_paragraph(ach.strip(), style='List Bullet')

    # Projects
    if projects:
        doc.add_heading('Projects', level=2)
        for proj in projects:
            p = doc.add_paragraph()
            p.add_run(proj.get('project_name', 'Unknown Project')).bold = True
            tech = proj.get('tech_stack', '')
            if tech:
                p.add_run(f" | {tech}")
            doc.add_paragraph(proj.get('description', ''))
            links = []
            if proj.get('github_url'):
                links.append(f"GitHub: {proj['github_url']}")
            if proj.get('live_url'):
                links.append(f"Live: {proj['live_url']}")
            if links:
                doc.add_paragraph(" - ".join(links))

    # Certifications
    if certs:
        doc.add_heading('Certifications', level=2)
        for cert in certs:
            p = doc.add_paragraph()
            p.add_run(cert.get('name', 'Unknown Cert')).bold = True
            issuer = cert.get('issuer', '')
            if issuer:
                p.add_run(f" | {issuer}")
            date = cert.get('issue_date')
            if date:
                p.add_run(f" | {date}")

    result_file = BytesIO()
    doc.save(result_file)
    return result_file.getvalue()
